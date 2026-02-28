# main.py
# BibleAI v281.38.Ωραία Εκκλησία (Orea Ekklisia) '아름다운교회'

# 수정사항:
# - [개선] 검색어 간격 20자로 축소 + 어구 검색 기능 추가 (정규식 기반)
# - [수정] 검색 결과 요약/복사/담기 버튼 작동 개선
# - [수정] 검색 히스토리 다시 검색 버튼 작동 수정
# - [추가] 외부 주석 모듈에도 요약/복사 버튼 추가
# - [개선] AI 요약 길이 동적 조정 (내용 길이에 비례)
# - [개선] 외부 주석 모듈 요약 기능 추가
# - [추가] Ollama 모델 선택 기능 (자동 감지 및 드롭다운)
# - [추가] 사이드바 프리셋 기본 접힘 상태로 시작
# - [개선] 크로스 플랫폼 호환성 향상
# - [FIXED] 6가지 사전 형식 모두 올바르게 표시되도록 수정
# - [FIXED] 막대바 담기 버튼, 담아줘 명령 후 바구니 오류 수정,  클로드 픽스 재수정 버전 
# - 2026-02-28 블로그 최종 배포본

import streamlit as st
import ollama
from groq import Groq
import os, re, sqlite3, warnings, glob, zlib
import fitz  # PyMuPDF
from docx import Document
from bs4 import BeautifulSoup
import ebooklib
from ebooklib import epub
from io import BytesIO
import streamlit.components.v1 as components
import platform
import subprocess
import shutil
import pyperclip  # 클립보드 복사를 위한 모듈
import tempfile
import webbrowser
client = Groq(api_key="")  # ← 여기에 본인의 Groq API 키를 입력하세요

from datetime import datetime
from collections import defaultdict
import json
from core.bible_utils import decode_rtf, get_ultimate_bible_map
from core.commentary_utils import load_commentaries_for_path, scan_commentary_files
from core.search_engine import parse_reference, fetch_intro, fetch_bible_text

warnings.filterwarnings('ignore')

# --- [NEW] 사전 파일 스캔 함수 ---
@st.cache_data(show_spinner=False)
def scan_dictionary_files():
    """dct 폴더 내의 모든 *.dct.twm 파일을 스캔합니다."""
    dct_folder = "dct"
    if not os.path.exists(dct_folder) or not os.path.isdir(dct_folder):
        return []
    
    files = []
    for filename in os.listdir(dct_folder):
        if filename.lower().endswith(".dct.twm"):
            full_path = os.path.join(dct_folder, filename)
            # 파일명에서 확장자 제거 (예: krstrong.dct.twm -> krstrong)
            display_name = filename[:-8] if filename.lower().endswith(".dct.twm") else filename
            files.append({
                "path": full_path,
                "filename": filename,
                "display_name": display_name
            })
    
    return sorted(files, key=lambda x: x['display_name'])

# --- [NEW] 사전 인덱스 컬럼 탐지 함수 ---
@st.cache_data(show_spinner=False)
def detect_dictionary_index_column(db_path):
    """사전 DB의 인덱스 컬럼명을 자동 탐지합니다."""
    if not os.path.exists(db_path):
        return None
    
    try:
        conn = sqlite3.connect(db_path)
        cur = conn.cursor()
        
        # topics 테이블의 컬럼 구조 확인
        cur.execute("PRAGMA table_info(topics)")
        columns = [col[1] for col in cur.fetchall()]
        conn.close()
        
        # 일반적인 인덱스 컬럼명 우선순위
        priority_columns = ['subject', 'topic', 'key', 'word', 'entry', 'term', 'id']
        
        for col in priority_columns:
            if col in columns:
                return col
        
        # 우선순위에 없으면 첫 번째 컬럼 반환
        return columns[0] if columns else None
        
    except Exception:
        return None

# --- [NEW] 사전에서 사용 가능한 인덱스 목록 가져오기 ---
@st.cache_data(show_spinner=False)
def get_dictionary_index_samples(db_path, index_column, limit=20):
    """사전에서 인덱스 샘플을 가져옵니다."""
    if not os.path.exists(db_path) or not index_column:
        return []
    
    try:
        conn = sqlite3.connect(db_path)
        cur = conn.cursor()
        cur.execute(f"SELECT DISTINCT {index_column} FROM topics LIMIT ?", (limit,))
        samples = [row[0] for row in cur.fetchall() if row[0]]
        conn.close()
        return samples
    except Exception:
        return []

# ===================================================================
# ★★★ 핵심 수정: 사전 검색 및 텍스트 정제 함수 (전면 재작성) ★★★
# 수정 이유:
#   - 기존 clean_lexicon_text_advanced()는 RTF \u 명령어를 삭제 후
#     16진수로 재해석하는 이중 오류로 한글·헬라어·히브리어가 깨짐
#   - BeautifulSoup str(soup) 변환 시 <br></p> 등 태그가 그대로 노출됨
# 수정 방향:
#   - core/bible_utils.py의 decode_rtf()를 핵심 디코딩에 사용 (10진수 \u 정확 처리)
#   - HTML 태그는 BeautifulSoup .get_text()로 텍스트만 추출
#   - zlib 압축 blob 지원 유지 (Mickelson's Strong 등)
# ===================================================================

@st.cache_data(show_spinner=False)
def get_lexicon_enhanced(db_path, search_term, index_column):
    """
    통합 사전 검색 함수 - 한글, 헬라어, 히브리어, 영어 모두 정상 출력
    1. 바이블렉스 (bible.dct.twm): RTF 유니코드 10진수 -> decode_rtf()
    2. Bullinger-App: 일반 텍스트
    3. Mickelson's Strong: zlib 압축 blob -> 압축해제 -> decode_rtf()
    4. 70인역, 킹제임스스트롱, 한글스트롱: RTF 유니코드 10진수 -> decode_rtf()
    """
    if not os.path.exists(db_path) or not index_column:
        return None, None

    try:
        conn = sqlite3.connect(db_path)
        cur = conn.cursor()

        # data, data2 모두 추출 시도
        try:
            query = f"SELECT c.data, c.data2 FROM content c JOIN topics t ON c.topic_id = t.id WHERE t.{index_column} = ? LIMIT 1"
            cur.execute(query, (search_term.strip(),))
            row = cur.fetchone()
        except Exception:
            # data2 컬럼 없는 사전은 data만 추출
            try:
                query = f"SELECT c.data FROM content c JOIN topics t ON c.topic_id = t.id WHERE t.{index_column} = ? LIMIT 1"
                cur.execute(query, (search_term.strip(),))
                r = cur.fetchone()
                row = (r[0], None) if r else None
            except Exception:
                row = None
        conn.close()

        if not row:
            return None, None

        raw_data, blob_data = row[0], row[1] if len(row) > 1 else None
        content = ""

        # === 1단계: zlib 압축 blob 처리 (Mickelson's Strong 등) ===
        if blob_data and isinstance(blob_data, bytes):
            try:
                decompressed = zlib.decompress(blob_data)
                content = decompressed.decode('utf-8', errors='ignore')
            except Exception:
                try:
                    content = blob_data.decode('utf-8', errors='ignore')
                except Exception:
                    content = str(blob_data)

        # === 2단계: 일반 data 처리 ===
        elif raw_data:
            if isinstance(raw_data, bytes):
                content = raw_data.decode('utf-8', errors='ignore')
            else:
                content = str(raw_data)

        if not content:
            return None, None

        # === 3단계: HTML 우선 처리 (CWSD 등 HTML 포함 사전) ===
        # decode_rtf() 전에 HTML을 먼저 처리해야 <grk><trn> 등 태그가 제거됨
        # decode_rtf() 이후에는 HTML이 이스케이프되어 태그 인식 불가
        if re.search(r'<[a-zA-Z/][^>]*>', content):
            try:
                soup = BeautifulSoup(content, 'html.parser')
                for tag in soup(['script', 'style']):
                    tag.decompose()
                for br in soup.find_all(['br', 'p', 'div', 'li', 'h1', 'h2', 'h3', 'h4']):
                    br.insert_before('\n')
                # 비표준 태그(<grk><trn><a class=T> 등) 포함 모든 태그 제거
                content = soup.get_text(separator='')
            except Exception:
                content = re.sub(r'<[^>]+>', '', content)

        # === 4단계: 핵심 디코딩 - decode_rtf() 사용 ===
        # RTF 유니코드 10진수(\u-숫자) 방식 정확 처리
        # 한글·헬라어·히브리어 모두 올바르게 변환됨
        plain_text = decode_rtf(content)

        # === 5단계: latin-1/cp1252 인코딩 깨짐 복구 (Lxx 등) ===
        if re.search(r'[\xe0-\xff]{2,}', plain_text):
            try:
                recovered = plain_text.encode('latin-1').decode('utf-8', errors='ignore')
                before_count = len(re.findall(r'[\u0370-\u03FF\u0590-\u05FF\uAC00-\uD7A3]', plain_text))
                after_count  = len(re.findall(r'[\u0370-\u03FF\u0590-\u05FF\uAC00-\uD7A3]', recovered))
                if after_count > before_count:
                    plain_text = recovered
            except Exception:
                pass

        # === 6단계: 공백 및 깨진 물음표 정리 ===
        plain_text = re.sub(r'([\u0590-\u05FF\u0370-\u03FF\uAC00-\uD7A3])\?', r'\1', plain_text)
        plain_text = re.sub(r'\n{3,}', '\n\n', plain_text)
        plain_text = plain_text.strip()

        # === 7단계: HTML 표시용 변환 ===
        html_content = plain_text.replace('\n', '<br>')

        return plain_text, html_content

    except Exception:
        return None, None

# ===================================================================
# ★★★ 수정 끝 ★★★
# ===================================================================


# --- 성경 검색 기능 정의 (파일 윗부분에 있어야 합니다) ---

# ========== [NEW] 6가지 개선사항 함수들 ==========

# [NEW 1] 관련도 점수 계산 함수
def calculate_relevance_score(content, search_query):
    """검색어와 문서의 관련도 점수를 계산합니다."""
    if not content or not search_query:
        return 0

    content_lower = content.lower()
    query_lower = search_query.lower()

    # 조건 검색 파싱
    words = query_lower.split()
    search_terms = [w[1:] if w.startswith(('+', '-')) else w for w in words if not w.startswith('-')]

    score = 0
    for term in search_terms:
        # 출현 빈도 점수
        count = content_lower.count(term)
        score += count * 10

        # 위치 점수 (앞부분에 나올수록 높은 점수)
        first_pos = content_lower.find(term)
        if first_pos != -1:
            if first_pos < len(content) * 0.1:
                score += 50
            elif first_pos < len(content) * 0.3:
                score += 30
            elif first_pos < len(content) * 0.5:
                score += 15

    return score

# [NEW 2] 검색 히스토리 관리 함수
def save_search_history(query):
    """검색어를 히스토리에 저장"""
    if 'search_history' not in st.session_state:
        st.session_state.search_history = []

    # 중복 제거하고 최근 항목을 앞에 추가
    if query in st.session_state.search_history:
        st.session_state.search_history.remove(query)

    st.session_state.search_history.insert(0, query)

    # 최대 20개까지만 저장
    st.session_state.search_history = st.session_state.search_history[:20]

# [NEW 5] AI 요약 기능 - 개선된 버전 (내용 길이에 따라 동적 조정)
def summarize_with_ai(content, model_name):
    """AI를 사용하여 검색 결과를 요약합니다. 내용 길이에 따라 요약 길이 조정."""
    try:
        content_length = len(content)

        # 내용 길이에 따라 요약 길이와 샘플링 범위 동적 조정
        if content_length < 500:
            summary_instruction = "2-3문장으로 핵심만 간단히"
            sample_length = content_length
        elif content_length < 1500:
            summary_instruction = "3-5문장으로 핵심 내용을"
            sample_length = 1500
        elif content_length < 3000:
            summary_instruction = "5-7문장으로 주요 내용을 상세히"
            sample_length = 3000
        else:
            summary_instruction = "7-10문장으로 전체 맥락과 주요 내용을 포괄적으로"
            sample_length = 5000

        prompt = f"""다음 내용을 {summary_instruction} 요약해주세요. 문맥이 연결되도록 작성하세요:

{content[:sample_length]}

요약:"""

        response = ollama.generate(model=model_name, prompt=prompt)
        return response['response']
    except Exception as e:
        return f"요약 실패: {str(e)}"

    # ========== [END NEW FUNCTIONS] ==========



# ========== [END NEW FUNCTIONS] ==========


def search_bible_sqlite(query):
    try:
        import sqlite3
        # 목사님의 성경 DB 파일 경로를 확인해주세요. 예: 'bible.db'
        conn = sqlite3.connect('bible_database.db')
        cursor = conn.cursor()

        # 성경 본문에서 검색 (테이블명과 컬럼명은 목사님 DB 설정에 맞춰야 합니다)
        sql = "SELECT book, chapter, verse, content FROM bible_table WHERE content LIKE ?"
        cursor.execute(sql, (f'%{query}%',))
        rows = cursor.fetchall()

        results = []
        for row in rows:
            results.append({
                'file': f"{row[0]} {row[1]}:{row[2]}",
                'content': row[3]
            })
        conn.close()
        return results
    except Exception as e:
        # DB가 없거나 설정이 다르면 빈 리스트 반환
        return []

# --- [NEW] 개선된 일반 파일 검색 기능 정의 ---
def search_files_advanced(query, selected_folders=None, include_extensions=None):
    """
    향상된 파일 검색 기능 v2.0
    - 조건 검색 지원: +필수단어, -제외단어
    - 어구 검색: 따옴표로 묶인 구문 정확히 검색 (예: "인간의 죄", '하나님의 사랑')
    - 주제어 검색: 컨텍스트 기반 스니펫 추출 (검색어 간격 20자 이내)
    - 다양한 파일 형식 지원

    Args:
        query: 검색어 (예: "+사랑 -미움", "인간의 죄", "'하나님의 사랑'")
        selected_folders: 검색할 폴더 리스트
        include_extensions: 검색할 파일 확장자 리스트
    """
    if selected_folders is None:
        selected_folders = ["."]

    if include_extensions is None:
        include_extensions = ['.txt', '.rtf', '.docx', '.pdf', '.epub', '.html', '.htm']

    # 어구 검색 파싱 (따옴표로 묶인 구문 추출)
    phrase_terms = []  # 어구 검색어 (따옴표로 묶인 것)

    # 쌍따옴표로 묶인 구문 추출
    double_quote_phrases = re.findall(r'"([^"]+)"', query)
    phrase_terms.extend(double_quote_phrases)

    # 홑따옴표로 묶인 구문 추출
    single_quote_phrases = re.findall(r"'([^']+)'", query)
    phrase_terms.extend(single_quote_phrases)

    # 따옴표로 묶인 부분을 제거한 나머지 쿼리
    remaining_query = re.sub(r'"[^"]+"', '', query)
    remaining_query = re.sub(r"'[^']+'", '', remaining_query)

    # 조건 검색 파싱
    include_terms = []  # +로 시작하는 필수 단어
    exclude_terms = []  # -로 시작하는 제외 단어
    normal_terms = []   # 일반 검색어

    words = remaining_query.split()
    for word in words:
        if not word.strip():
            continue
        if word.startswith('+'):
            include_terms.append(word[1:].lower())
        elif word.startswith('-'):
            exclude_terms.append(word[1:].lower())
        else:
            normal_terms.append(word.lower())

    # 모든 검색어를 하나로 통합 (일반 검색어 + 필수 단어)
    all_search_terms = normal_terms + include_terms + [p.lower() for p in phrase_terms]

    results = []

    # 파일 수집
    files_to_search = []
    for folder in selected_folders:
        folder_path = folder if folder != "." else "."
        if os.path.exists(folder_path) and os.path.isdir(folder_path):
            for root, dirs, files in os.walk(folder_path):
                for file in files:
                    if any(file.lower().endswith(ext) for ext in include_extensions):
                        if not file.startswith('.'):
                            files_to_search.append(os.path.join(root, file))

    # 파일 검색 수행
    for file_path in files_to_search:
        try:
            content = read_file(file_path)
            if not content:
                continue

            content_lower = content.lower()

            # 제외 단어 체크
            should_skip = False
            for exclude_term in exclude_terms:
                if exclude_term in content_lower:
                    should_skip = True
                    break

            if should_skip:
                continue

            # 어구 검색 체크 (정확한 매칭)
            if phrase_terms:
                has_all_phrases = True
                for phrase in phrase_terms:
                    if phrase.lower() not in content_lower:
                        has_all_phrases = False
                        break

                if not has_all_phrases:
                    continue

            # 필수 단어 체크 (모두 포함되어야 함)
            has_all_required = True
            for include_term in include_terms:
                if include_term not in content_lower:
                    has_all_required = False
                    break

            if not has_all_required:
                continue

            # 일반 검색어 체크 (하나라도 포함되면 OK)
            if normal_terms:
                has_any_normal = False
                for normal_term in normal_terms:
                    if normal_term in content_lower:
                        has_any_normal = True
                        break

                if not has_any_normal:
                    continue

            # 결과 생성 - 컨텍스트 기반 스니펫 추출 (20자 간격)
            if all_search_terms or phrase_terms:
                # 어구 검색 먼저 확인
                first_match_pos = -1
                matched_term = ""

                # 어구 검색 우선
                for phrase in phrase_terms:
                    pos = content.lower().find(phrase.lower())
                    if pos != -1:
                        if first_match_pos == -1 or pos < first_match_pos:
                            first_match_pos = pos
                            matched_term = phrase

                # 어구 검색에서 못 찾았으면 개별 검색어로
                if first_match_pos == -1:
                    for term in all_search_terms:
                        pos = content_lower.find(term)
                        if pos != -1:
                            if first_match_pos == -1 or pos < first_match_pos:
                                first_match_pos = pos
                                matched_term = term

                if first_match_pos != -1:
                    # 검색어 간격 검증 (복합 검색어의 경우 20자 이내에 모두 있어야 함)
                    if len(all_search_terms) > 1 or len(phrase_terms) > 0:
                        # 20자 윈도우 내에서 모든 검색어 확인
                        window_size = 20
                        all_terms_to_check = all_search_terms + [p.lower() for p in phrase_terms]

                        # 첫 검색어 기준으로 윈도우 설정
                        window_start = max(0, first_match_pos - window_size)
                        window_end = min(len(content), first_match_pos + len(matched_term) + window_size)
                        window_content = content[window_start:window_end].lower()

                        # 윈도우 내에 모든 검색어가 있는지 확인
                        all_found_in_window = all(term in window_content for term in all_terms_to_check if term != matched_term)

                        if not all_found_in_window and len(all_terms_to_check) > 1:
                            # 검색어가 20자 이내에 없으면 스킵
                            continue

                    # 앞뒤 200자씩 추출 (컨텍스트 포함)
                    start = max(0, first_match_pos - 200)
                    end = min(len(content), first_match_pos + len(matched_term) + 200)
                    snippet = content[start:end]

                    # 앞뒤 생략 표시
                    if start > 0:
                        snippet = "..." + snippet
                    if end < len(content):
                        snippet = snippet + "..."

                    # 모든 매칭 위치 찾기
                    all_matches = []
                    for term in all_search_terms:
                        pos = 0
                        while True:
                            pos = content_lower.find(term, pos)
                            if pos == -1:
                                break
                            all_matches.append(pos)
                            pos += 1

                    # 매칭 개수 표시
                    match_count = len(all_matches)

                    # [NEW 1] 관련도 점수 계산
                    relevance_score = calculate_relevance_score(content, query)

                    results.append({
                        'file': os.path.basename(file_path),
                        'content': f"[검색어 '{query}' - {match_count}건 발견]\n\n{snippet}\n\n--- 전체 내용 ---\n{content}",
                        'relevance_score': relevance_score
                    })
            else:
                # 검색어 없이 조건만 있는 경우 (예: "-인내")
                # [NEW 1] 관련도 점수 계산
                relevance_score = calculate_relevance_score(content, query)

                results.append({
                    'file': os.path.basename(file_path),
                    'content': content,
                    'relevance_score': relevance_score
                })

        except Exception as e:
            continue

    # [NEW 1] 관련도 점수로 정렬 (내림차순)
    results.sort(key=lambda x: x.get('relevance_score', 0), reverse=True)

    return results

# --- [1. 시스템 설정 및 세션 초기화] ---
st.set_page_config(page_title="Ωραία Εκκλησία (Orea Ekklisia) '아름다운교회'", layout="wide")

# --- [조립용 함수: 설정값들을 텍스트로 합침] ---
def get_custom_prompt(context_data):
    # 세션 보관함에서 아까 선택한 10가지 값을 쏙쏙 뽑아옵니다.
    c1, c2, c3 = st.session_state['c1'], st.session_state['c2'], st.session_state['c3']
    c4, c5, c6 = st.session_state['c4'], st.session_state['c5'], st.session_state['c6']
    c7, c8, c9 = st.session_state['c7'], st.session_state['c8'], st.session_state['c9']
    c10 = st.session_state['c10']

    return f"""당신은 {c1} 전통을 따르는 숙련된 강해설교가입니다.
제공된 자료를 바탕으로 성령 충만한 설교 초안을 작성하세요.

[설교 기획 가이드라인]
1. 신학적 노선: {c1} 관점
2. 예배 상황: {c6} ({c4} 대상 / {c5} 분량)
3. 본문 번역: {c8} 기준으로 인용
4. 구조 및 문체: {c7} 구조, {c2} 스타일로 {c3} 작성
5. 적용 강조점: {c9}에 중점을 둠
6. 특별 요청: {c10 if c10 else "본문 중심의 깊이 있는 주해와 실제적 삶의 적용"}

[참고 자료들]:
{context_data}

---
위 기획안과 자료를 바탕으로 은혜로운 설교 초안을 작성하라.""".strip()

THEWORD_DB = "bible.dct.twm" # 사전용 DB

# [NEW 2, 3] 검색 히스토리와 바구니 그룹 초기화 추가
keys = {
    'basket': [],
    'scan_res': [],
    'v_content': '',
    'current_path': os.getcwd(),
    'show_lex': True,
    'selected_model': 'gemma3:4b',
    'search_history': [],  # [NEW 2] 검색 히스토리
    'basket_groups': defaultdict(list),  # [NEW 3] 바구니 그룹
    'current_group': "기본 그룹",  # [NEW 3] 현재 선택된 그룹
    'trigger_search': None,  # [NEW] 검색 히스토리 트리거
    'ai_response': '',  # AI 응답 저장
    'last_user_input': '',  # 마지막 사용자 입력 저장
    'show_ai_response': False,  # AI 응답 표시 여부
}

for key, default in keys.items():
    if key not in st.session_state:
        if key == 'basket_groups':
            st.session_state[key] = defaultdict(list)
        else:
            st.session_state[key] = default

# --- [1.5 상시 표시 설정 패널: 9-Options] ---
# --- [1.5 프리셋 기능이 추가된 설정 패널] ---
with st.sidebar:
    st.header("🎚️ 설교 믹서 프리셋")

    # 1. 프리셋 정의 (믹서기 설정값 세트)
    presets = {
        "Manual (직접 설정)": None,
        "☀️ 주일 대예배": {
            'c1': "개혁적", 'c2': "학술적/깊이있는", 'c3': "엄밀하게", 'c4': "장년",
            'c5': "40분", 'c6': "주일대예배", 'c7': "3대지(대지형)", 'c8': "개역개정", 'c9': "교리/신학적 지식"
        },
        "🌅 새벽 기도의 등불": {
            'c1': "복음주의적", 'c2': "대중적/쉬운", 'c3': "따뜻하게", 'c4': "장년",
            'c5': "10분", 'c6': "새벽기도", 'c7': "강해중심", 'c8': "개역개정", 'c9': "개인적 위로/회복"
        },
        "🔥 금요 철야 부흥": {
            'c1': "오순절적", 'c2': "대중적/쉬운", 'c3': "웅변적으로", 'c4': "장년",
            'c5': "20분", 'c6': "금요철야", 'c7': "예화/간증 중심", 'c8': "개역개정", 'c9': "사명/전도"
        },
        "☕ 청년부 바이블톡": {
            'c1': "복음주의적", 'c2': "대중적/쉬운", 'c3': "대화형으로", 'c4': "청년",
            'c5': "20분", 'c6': "주일저녁예배", 'c7': "원포인트", 'c8': "새번역", 'c9': "제자도/사회적 책임"
        }
    }

    # 2. 프리셋 선택창
    selected_preset = st.selectbox("🎯 상황별 프리셋 선택", list(presets.keys()))

    # 3. 프리셋 적용 로직
    if selected_preset != "Manual (직접 설정)":
        preset_values = presets[selected_preset]
        for key, value in preset_values.items():
            st.session_state[key] = value

    st.markdown("---")
    st.header("🎨 세부 옵션 조정")

    # 이후 기존의 st.expander 및 selectbox 코드들이 그대로 옵니다.
    # (이미 index=...index(st.session_state['c1']) 로직을 짜두었기 때문에
    # 프리셋을 바꾸면 아래 화면도 자동으로 바뀝니다!)

    st.header("🎨 설교 프롬프트 설정")

    # 세션 상태 초기화 (값이 없을 경우를 대비해 기본값 설정)
    opts = {
        'c1': "개혁적", 'c2': "학술적/깊이있는", 'c3': "엄밀하게", 'c4': "장년",
        'c5': "40분", 'c6': "주일대예배", 'c7': "3대지(대지형)",
        'c8': "개역개정", 'c9': "개인적 위로/회복", 'c10': ""
    }
    for k, v in opts.items():
        if k not in st.session_state: st.session_state[k] = v

    # 1. 예배 및 대상
    with st.expander("📍 1. 예배 및 대상", expanded=False):
        st.session_state['c1'] = st.selectbox("신학적 배경", ["개혁적", "웨슬리안", "오순절적", "침례교적", "복음주의적"], index=["개혁적", "웨슬리안", "오순절적", "침례교적", "복음주의적"].index(st.session_state['c1']))
        st.session_state['c4'] = st.selectbox("청중 대상", ["장년", "청년", "청소년", "아동"], index=["장년", "청년", "청소년", "아동"].index(st.session_state['c4']))
        st.session_state['c6'] = st.selectbox("예배 성격", ["주일대예배", "주일저녁예배", "수요예배", "금요철야", "새벽기도"], index=["주일대예배", "주일저녁예배", "수요예배", "금요철야", "새벽기도"].index(st.session_state['c6']))
        st.session_state['c5'] = st.selectbox("예상 소요 시간", ["1시간", "40분", "20분", "10분", "5분"], index=["1시간", "40분", "20분", "10분", "5분"].index(st.session_state['c5']))

    # 2. 구성 및 문체
    with st.expander("✍️ 2. 구성 및 문체", expanded=False):
        st.session_state['c2'] = st.selectbox("학술 수준", ["학술적/깊이있는", "대중적/쉬운"], index=["학술적/깊이있는", "대중적/쉬운"].index(st.session_state['c2']))
        st.session_state['c3'] = st.selectbox("어조(Tone)", ["엄밀하게", "따뜻하게", "웅변적으로", "대화형으로"], index=["엄밀하게", "따뜻하게", "웅변적으로", "대화형으로"].index(st.session_state['c3']))
        st.session_state['c7'] = st.selectbox("설교 구조", ["3대지(대지형)", "원포인트", "강해중심", "예화/간증 중심"], index=["3대지(대지형)", "원포인트", "강해중심", "예화/간증 중심"].index(st.session_state['c7']))
        st.session_state['c8'] = st.selectbox("성경 번역본", ["개역개정", "새번역", "쉬운성경", "NIV", "ESV"], index=["개역개정", "새번역", "쉬운성경", "NIV", "ESV"].index(st.session_state['c8']))

    # 3. 적용 강조점
    with st.expander("🎯 3. 적용 강조점", expanded=False):
        st.session_state['c9'] = st.selectbox("적용 방향", ["개인적 위로/회복", "사명/전도", "제자도/사회적 책임", "교리/신학적 지식"], index=["개인적 위로/회복", "사명/전도", "제자도/사회적 책임", "교리/신학적 지식"].index(st.session_state['c9']))
        st.session_state['c10'] = st.text_area("✨ 특별 요청사항", value=st.session_state['c10'], placeholder="예: 고난 받는 성도들에게 소망을...", height=100)

    st.caption("💡 모든 설정은 '초안 생성' 시 자동으로 반영됩니다.")
    st.markdown("---")

# --- [수정] AI 엔진 선택 및 자동 인식 기능 ---
with st.sidebar:
    st.header("🤖 AI 엔진 설정")
    
    # 1. 엔진 선택 (기본 Groq 권장)
    engine_choice = st.radio(
        "사용할 엔진을 선택하세요:",
        ["Groq (온라인/초고속)", "Ollama (로컬/서버)"],
        index=0,
        help="Groq 토큰 소진 시 Ollama로 전환하여 사용하세요."
    )

    if engine_choice == "Groq (온라인/초고속)":
        st.info("⚡ Groq API Key 넣으셨나요? 꼭 확인해 주세요~")
        st.session_state.selected_engine = "groq"
        st.session_state.selected_model = "llama-3.3-70b-versatile"
    
    else:
        st.session_state.selected_engine = "ollama"
        try:
            # 2. Ollama 설치 모델 목록 자동 인식
            import ollama
            response = ollama.list()
            models = [m.get('name') or m.get('model') for m in response.get('models', [])]
            
            if models:
                # 목록이 있으면 선택창 표시
                selected_ollama = st.selectbox("사용할 Ollama 모델 선택:", options=models)
                st.session_state.selected_model = selected_ollama
                st.success(f"✅ {selected_ollama} 준비 완료")
            else:
                # 모델이 없을 경우 온라인(Groq)으로 자동 제안
                st.warning("⚠️ 설치된 모델이 없어 온라인 모드를 권장합니다.")
                st.session_state.selected_engine = "groq"
                st.session_state.selected_model = "llama-3.3-70b-versatile"
                
        except Exception:
            # Ollama 서버 미실행 시 온라인으로 자동 우회 유도
            st.error("⚠️ Ollama 서버 연결 실패")
            st.info("온라인(Groq) 엔진을 사용하거나, PC에서 Ollama 앱을 켜주세요.")
            st.session_state.selected_engine = "groq"
            st.session_state.selected_model = "llama-3.3-70b-versatile"

# 이 아랫줄에 바로 '# [NEW 2] 검색 히스토리 사이드바 추가'가 오게 됩니다.
# [NEW 2] 검색 히스토리 사이드바 추가
with st.sidebar:
    st.markdown("---")
    st.header("📜 최근 검색어")

    if st.session_state.search_history:
        selected_history = st.selectbox(
            "검색 히스토리",
            [""] + st.session_state.search_history,
            format_func=lambda x: "선택하세요..." if x == "" else f"🔍 {x}",
            key="history_selector"
        )

        if selected_history and st.button("🔄 다시 검색", use_container_width=True, key="rerun_search"):
            # 검색 실행 트리거 설정
            st.session_state.trigger_search = selected_history
    else:
        st.caption("검색 기록이 없습니다")

# [NEW 3] 바구니 그룹 관리 사이드바 추가
with st.sidebar:
    st.markdown("---")
    st.header("📁 바구니 그룹 관리")

    # 그룹 목록 표시
    all_groups = list(set(list(st.session_state.basket_groups.keys()) + ["기본 그룹"]))

    col1, col2 = st.columns([3, 1])
    with col1:
        current_group = st.selectbox(
            "현재 작업 그룹",
            all_groups,
            index=all_groups.index(st.session_state.current_group) if st.session_state.current_group in all_groups else 0,
            key="group_selector"
        )
        st.session_state.current_group = current_group

    with col2:
        if st.button("➕", help="새 그룹 추가", key="add_group_btn"):
            st.session_state.show_new_group = True

    # 새 그룹 추가 UI
    if st.session_state.get('show_new_group', False):
        new_group_name = st.text_input("새 그룹 이름", key="new_group_input")
        col_a, col_b = st.columns(2)
        with col_a:
            if st.button("생성", key="create_group_btn"):
                if new_group_name:
                    st.session_state.basket_groups[new_group_name] = []
                    st.session_state.current_group = new_group_name
                    st.session_state.show_new_group = False
        with col_b:
            if st.button("취소", key="cancel_group_btn"):
                st.session_state.show_new_group = False

    # 현재 그룹의 항목 수 표시
    group_items = st.session_state.basket_groups.get(st.session_state.current_group, [])
    st.caption(f"📊 {st.session_state.current_group}: {len(group_items)}개 항목")

# --- [2. 엔진 로직: 파일 읽기 및 RTF 디코드] ---

@st.cache_data(show_spinner=False)
def read_file(path):
    if not os.path.exists(path): return ""
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".docx":
            doc = Document(path)
            return "\n".join([p.text for p in doc.paragraphs])
        elif ext == ".pdf":
            text = ""
            with fitz.open(path) as doc:
                for page in doc: text += page.get_text()
            return text
        elif ext in [".txt", ".rtf"]:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
                return decode_rtf(content) if ext == ".rtf" else content
        elif ext == ".epub":
            book = epub.read_epub(path)
            items = []
            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    soup = BeautifulSoup(item.get_content(), 'html.parser')
                    for link in soup.find_all('a'):
                        link.unwrap()
                    items.append(soup.get_text())
            return "\n".join(items)
        elif ext in ['.html', '.htm']:
            with open(path, 'r', encoding='utf-8') as f:
                return BeautifulSoup(f.read(), 'html.parser').get_text()
    except Exception as e:
        return f"파일 읽기 오류 ({path}): {str(e)}"
    return ""

# --- [3. 성경 지명/약어 매핑] ---
BIBLE_ALIAS_FLAT, BIBLE_RAW_MAP = get_ultimate_bible_map()
# --- [4. 검색 및 외부 주석 엔진] ---
@st.cache_data(show_spinner=False)
def search_engine(text, user_book, chap, verse_input):
    """
    Bible 텍스트에서 서론/본문을 검색하는 라우터 함수입니다.
    - parse_reference: 입력 파싱 및 모드 결정
    - fetch_intro: 책/장 서론 추출
    - fetch_bible_text: 일반 절 본문 추출
    """
    parsed = parse_reference(user_book, chap, verse_input, BIBLE_ALIAS_FLAT, BIBLE_RAW_MAP)
    if not parsed:
        return ""

    std, norm_chap, verses, mode = parsed
    results_dict = {}

    # 서론(책/장) 처리
    intro_results = fetch_intro(
        text,
        std,
        norm_chap,
        verse_input,
        BIBLE_ALIAS_FLAT,
        BIBLE_RAW_MAP,
    )
    results_dict.update(intro_results)

    # 일반 절 본문 처리
    if mode == "verse":
        bible_results = fetch_bible_text(
            text,
            std,
            norm_chap,
            verses,
            BIBLE_ALIAS_FLAT,
            BIBLE_RAW_MAP,
        )
        results_dict.update(bible_results)

    res_list = [f"{key}\n{content}" for key, content in results_dict.items()]
    return "\n\n".join(res_list) if res_list else ""

def get_external_commentaries(user_book, chap, vers, selected_folders=None):
    """
    외부 주석/성경 DB 파일(.mybible, .twm, .sqlite3, .cdb 등)을 모두 스캔한 뒤,
    각 파일 형식별 로더(core.commentary_utils)를 통해 주석을 통합합니다.
    """
    if selected_folders is None:
        selected_folders = ["."]

    # 1) 검색 대상 파일 수집 (동적 스캔)
    com_files = scan_commentary_files(selected_folders)

    # 2) 성경 책 이름을 표준 book_id로 변환
    bible_std_list = list(BIBLE_RAW_MAP.keys())
    try:
        normalized_book = user_book.strip()
        book_match = re.match(r"^([가-힣a-zA-Z0-9]+)", normalized_book)
        if book_match:
            book_part = book_match.group(1)
            std_name = BIBLE_ALIAS_FLAT.get(book_part.lower())
        else:
            std_name = BIBLE_ALIAS_FLAT.get(normalized_book.lower())

        std_name_upper = std_name.upper() if std_name else None
        book_id = None
        for i, book in enumerate(bible_std_list):
            if book.upper() == std_name_upper:
                book_id = i + 1
                break

        if book_id is None:
            raise ValueError(f"Book not found in list: {std_name}")
    except (ValueError, TypeError):
        return ""

    # 3) 파일별 로더를 호출해서 결과 통합
    com_results = []
    chap_int = int(chap)
    vers_int = int(vers)

    for path in com_files:
        com_results.extend(load_commentaries_for_path(path, int(book_id), chap_int, vers_int))

    # 4) 중복 제거
    seen = set()
    unique_com_results = []
    for result in com_results:
        if result not in seen:
            seen.add(result)
            unique_com_results.append(result)

    return "\n\n".join(unique_com_results)

@st.cache_data(show_spinner=False)
def get_lexicon(code):
    if not os.path.exists(THEWORD_DB): return None
    try:
        conn = sqlite3.connect(THEWORD_DB); cur = conn.cursor()
        cur.execute("SELECT c.data FROM content c JOIN topics t ON c.topic_id = t.id WHERE t.subject = ? LIMIT 1", (code.upper().strip(),))
        row = cur.fetchone(); conn.close()
        return decode_rtf(row[0]) if row else None
    except: return None

# --- [5. UI 레이아웃] ---
with st.sidebar:
    st.title("🎂 v281.36.Ωραία Εκκλησία (Orea Ekklisia) '아름다운교회'")

    # 수정된 폴더 정렬 로직
    def sort_folders(folders):
        """폴더를 이름순으로 정렬하되 숫자로 시작하는 폴더는 맨 앞으로 배치"""
        numeric_first = []
        alpha_rest = []

        for folder in folders:
            if folder and folder[0].isdigit():
                numeric_first.append(folder)
            else:
                alpha_rest.append(folder)

        # 숫자로 시작하는 폴더는 숫자 순으로 정렬
        numeric_first.sort(key=lambda x: [int(c) if c.isdigit() else c for c in re.split('([0-9]+)', x)])
        # 알파벳 폴더는 일반 정렬
        alpha_rest.sort()

        return numeric_first + alpha_rest

    # 모든 폴더 가져오기 및 정렬
    all_dirs = [d for d in os.listdir('.') if os.path.isdir(d) and not d.startswith('.')]
    sorted_dirs = sort_folders(all_dirs)

    # docs 폴더가 있다면 기본 선택
    default_selection = ["docs"] if "docs" in sorted_dirs else []

    # 검색 범위 선택 (정렬된 순서대로 표시)
    selected_folders = st.multiselect(
        "📍 검색 범위(폴더)",
        sorted_dirs,
        default=default_selection
    )

    st.divider()

    if 'show_lex' not in st.session_state:
        st.session_state.show_lex = True

    show_lex_current = st.checkbox("🏺 원어 해설창 토글", value=st.session_state.show_lex, key="original_lex_toggle")
    if show_lex_current != st.session_state.show_lex:
        st.session_state.show_lex = show_lex_current

    if st.session_state.show_lex:
        # [NEW] 사전 파일 스캔 및 선택
        dict_files = scan_dictionary_files()
        
        if dict_files:
            # 사전 파일 드롭다운
            dict_options = {f"{d['display_name']}": d for d in dict_files}
            selected_dict_name = st.selectbox(
                "📚 사전 선택",
                list(dict_options.keys()),
                key="dict_selector"
            )
            
            selected_dict_info = dict_options[selected_dict_name]
            dict_path = selected_dict_info['path']

            # [NEW] 인덱스 컬럼 자동 탐지
            index_column = detect_dictionary_index_column(dict_path)
            
            if index_column:
                st.caption(f"🔍 검색 기준: {index_column}")
                
                # [NEW] 인덱스 샘플 가져오기
                index_samples = get_dictionary_index_samples(dict_path, index_column, limit=20)
                
                if index_samples:
                    st.caption(f"💡 예시: {', '.join(index_samples[:5])}")
                
                # --- [수정] 사전 유형에 따른 초기값 설정 ---
                is_bullinger = "bullinger" in selected_dict_name.lower()
                # Bullinger는 숫자 위주, 나머지는 스트롱코드(G26) 위주
                default_val = "1" if is_bullinger else ("G26" if index_column.lower() == "subject" else "")
                
                search_term = st.text_input(
                    f"{index_column} 입력",
                    value=default_val,
                    key="lexicon_search"
                )
                
                if search_term:
                    # [개선] 압축 해제 및 유니코드 복원이 포함된 함수 호출
                    plain_text, html_content = get_lexicon_enhanced(dict_path, search_term, index_column)
                    
                    if html_content:
                        lex_h = st.slider("해설창 높이", 200, 800, 400, key="lex_height")
                        st.markdown("### 📖 원어 해설")
                        
                        # [개선] HTML 리치 텍스트 렌더링 (디자인 강화)
                        html_display = f"""
                        <div style="
                            background: linear-gradient(135deg, #fffcfc 0%, #f9f2f2 100%);
                            padding: 20px;
                            border-radius: 12px;
                            border: 2px solid #ff6b6b;
                            height: {lex_h}px;
                            overflow-y: auto;
                            font-family: 'Noto Serif KR', 'Malgun Gothic', serif;
                            line-height: 1.8;
                            color: #222;
                        ">
                            <h3 style="color: #d32f2f; margin-top: 0; font-size: 1.1em;">
                                📚 {search_term} - {selected_dict_name}
                            </h3>
                            <hr style="border: 1px solid #ff6b6b; margin: 15px 0; opacity: 0.3;">
                            <div style="font-size: 15px; word-break: break-all;">
                                {html_content}
                            </div>
                        </div>
                        """
                        st.markdown(html_display, unsafe_allow_html=True)
                    else:
                        st.warning(f"'{search_term}'에 대한 결과를 찾을 수 없습니다.")
            else:
                st.error("사전 구조를 인식할 수 없습니다.")
        else:
            st.info("💡 'dct' 폴더에 *.dct.twm 파일을 넣어주세요.")

    st.divider()

    # 동적 스캔된 주석 파일 목록 표시
    commentary_files = scan_commentary_files(selected_folders if selected_folders else ["."])
    st.caption(f"현재 인식된 주석 파일: {len(commentary_files)}개")

    if commentary_files:
        with st.expander(f"📚 외부 주석 파일 ({len(commentary_files)}개)", expanded=False):
            for cf in commentary_files:
                st.text(f"• {cf}")

    if st.button("🧹 전체 화면 지우기", use_container_width=True):
        st.session_state.scan_res = []
        st.session_state.v_content = ""

    st.subheader(f"🧺 바구니 ({len(st.session_state.basket)}개)")
    if st.button("🗑️ 바구니 비우기"): st.session_state.basket = []

    if st.session_state.basket:
        if st.button("🤖 LLM 통합 질문 생성"):
            context = "\n\n".join([i['content'] for i in st.session_state.basket])
            st.session_state.v_content = f"당신은 세계적인 신학자이자 성경언어학자입니다. 다음에 제시된 내용에 근거하여 상세히 설명하시오.\n\n{context}"

        doc = Document()
        doc.add_heading("Bible Research Report", 0)
        for item in st.session_state.basket:
            content_parts = item['content'].split('\n', 1)
            if len(content_parts) >= 2:
                title_line = content_parts[0].strip()
                content_body = content_parts[1].strip()
                clean_title = title_line.replace('####', '').strip()

                doc.add_heading(f"Source File: {item['file']}", level=0)
                if clean_title.startswith('[') and clean_title.endswith(']'):
                    clean_title_without_brackets = clean_title[1:-1]
                    enhanced_title = f"[{clean_title_without_brackets} - {item['file']}]"
                else:
                    enhanced_title = clean_title

                doc.add_heading(enhanced_title, level=1)
                doc.add_paragraph(content_body)
            else:
                doc.add_heading(f"Source File: {item['file']}", level=0)
                doc.add_paragraph(item['content'])
        bio = BytesIO(); doc.save(bio)
        st.download_button("📝 연구보고서(.docx) 저장", data=bio.getvalue(), file_name="BibleAI_Report.docx", use_container_width=True)

col_l, col_r = st.columns([0.45, 0.55])

with col_l:
    st.title("⚔️Ωραία Εκκλησία (Orea Ekklisia) '아름다운교회'")
    t1, t2, t3 = st.tabs(["📖 보화 찾기", "📁 서재 관리", "📝 프롬프트"])

    with t1:
        # [NEW 6] 바구니 편집기 추가
        with st.expander(f"🧺 바구니 편집기 ({len(st.session_state.basket)}건)", expanded=False):
            if st.session_state.basket:
                st.markdown("### 📝 바구니 내용 편집")

                # 편집 가능한 텍스트 영역
                basket_text = "\n\n" + "="*50 + "\n\n".join([
                    f"📄 {item['file']}\n{'-'*50}\n{item['content']}"
                    for item in st.session_state.basket
                ])

                edited_text = st.text_area(
                    "바구니 내용을 직접 수정할 수 있습니다",
                    value=basket_text,
                    height=400,
                    help="내용을 자유롭게 편집하세요. 저장 버튼을 누르면 변경사항이 적용됩니다.",
                    key="basket_editor"
                )

                col1, col2, col3 = st.columns(3)

                with col1:
                    if st.button("💾 편집 내용 저장", use_container_width=True, key="save_edit"):
                        st.session_state.basket = [{
                            'file': f"편집됨_{datetime.now().strftime('%H%M%S')}",
                            'content': edited_text
                        }]
                        st.success("✅ 편집 내용이 저장되었습니다!")

                with col2:
                    if st.button("📋 전체 복사", use_container_width=True, key="copy_basket_all"):
                        try:
                            pyperclip.copy(basket_text)
                            st.success("✅ 클립보드에 복사되었습니다!")
                        except Exception as e:
                            st.error(f"❌ 복사 실패: {e}")

                with col3:
                    if st.button("🗑️ 전체 삭제", use_container_width=True, key="delete_basket_all"):
                        st.session_state.basket = []
                        st.success("✅ 바구니가 비워졌습니다!")
            else:
                st.info("바구니가 비어 있습니다. 검색 후 자료를 담아주세요.")

        st.markdown("---")

        c1, c2, c3 = st.columns([2,1,1])
        b_in, ch_in, vs_in = c1.text_input("성경", "요"), c2.text_input("장", "6"), c3.text_input("절", "26-27")

        parsed_book, parsed_chap, parsed_vs = (None, None, None)
        if parsed_book and parsed_chap and parsed_vs:
            actual_book = parsed_book
            actual_chap = parsed_chap
            actual_vs = parsed_vs
        else:
            actual_book = b_in
            actual_chap = ch_in
            actual_vs = vs_in

        if st.button("🔍 전수 조사 시작", use_container_width=True):
            st.session_state.scan_res = []
            scan_dirs = selected_folders if selected_folders else ["."]
            files = [os.path.join(r, f) for d in scan_dirs for r, _, fs in os.walk(d) for f in fs if not f.startswith('.')]
            prog = st.progress(0); stat = st.empty()

            normalized_book = actual_book.strip()
            for i, p in enumerate(files):
                stat.text(f"탐색 중: {os.path.basename(p)}")
                res = search_engine(read_file(p), normalized_book, actual_chap, actual_vs)
                if res:
                    content_lines = res.split('\n', 1)
                    if len(content_lines) >= 2 and content_lines[0].startswith('#### ['):
                        bible_ref = content_lines[0].replace('#### ', '')
                        st.session_state.scan_res.append({"file": os.path.basename(p), "content": f"{bible_ref}\n{content_lines[1]}"})
                    else:
                        st.session_state.scan_res.append({"file": os.path.basename(p), "content": res})
                prog.progress((i+1)/len(files))

            verses_to_search = []
            if "-" in actual_vs:
                try:
                    start, end = map(int, actual_vs.split("-"))
                    verses_to_search = list(range(start, end + 1))
                except:
                    verses_to_search = [int(actual_vs)] if actual_vs.isdigit() else []
            elif "/" in actual_vs:
                try:
                    verses_to_search = [int(v.strip()) for v in actual_vs.split("/") if v.strip().isdigit()]
                except:
                    verses_to_search = []
            else:
                try:
                    verses_to_search = [int(actual_vs)]
                except:
                    verses_to_search = []

            if verses_to_search:
                stat.text(f"외부 주석 검색 중... ({len(verses_to_search)}개 절: {', '.join(map(str, verses_to_search))})")

                for i, verse_num in enumerate(verses_to_search):
                    stat.text(f"외부 주석 검색 중... (절 {verse_num} 처리 중 {i+1}/{len(verses_to_search)})")
                    comm_res = get_external_commentaries(normalized_book, int(actual_chap), verse_num, selected_folders)
                    if comm_res:
                        comm_sections = comm_res.split("\n\n#### 📚 [")
                        for idx, section in enumerate(comm_sections):
                            if idx == 0:
                                if section.strip():
                                    lines = section.split('\n', 1)
                                    if len(lines) > 1 and lines[0].startswith('#### 📚 ['):
                                        file_title = lines[0].replace("#### 📚 [", "").replace("]", "")
                                        content = lines[1]
                                        import re
                                        bible_ref_match = re.search(r'\[(.*?)\]', content)
                                        if bible_ref_match:
                                            bible_ref = bible_ref_match.group(0)
                                            content_without_filename = re.sub(r'\[(.*?)\]', '', content, 1).strip()
                                            st.session_state.scan_res.append({"file": f"📚 {file_title}", "content": f"{bible_ref}\n{content_without_filename}"})
                                        else:
                                            st.session_state.scan_res.append({"file": f"📚 {file_title}", "content": f"#### 📚 [{file_title}]\n{content}"})
                                    else:
                                        st.session_state.scan_res.append({"file": "📚 외부 주석 모듈", "content": section})
                            else:
                                full_section = "[" + section
                                lines = full_section.split('\n', 1)
                                if len(lines) > 1:
                                    file_title = lines[0].replace("#### 📚 [", "").replace("]", "")
                                    content = lines[1]
                                    import re
                                    bible_ref_match = re.search(r'\[(.*?)\]', content)
                                    if bible_ref_match:
                                        bible_ref = bible_ref_match.group(0)
                                        content_without_filename = re.sub(r'\[(.*?)\]', '', content, 1).strip()
                                        st.session_state.scan_res.append({"file": f"📚 {file_title}", "content": f"{bible_ref}\n{content_without_filename}"})
                                    else:
                                        st.session_state.scan_res.append({"file": f"📚 {file_title}", "content": f"#### 📚 [{file_title}]\n{content}"})

                stat.text(f"외부 주석 검색 완료! {len(st.session_state.scan_res)}개 결과")
            else:
                stat.text(f"검색 완료! {len(st.session_state.scan_res)}개 결과")

            stat.success(f"🎊 {len(st.session_state.scan_res)}개 발견!")

    with t2:
        st.markdown(f"📍 `{st.session_state.current_path}`")
        if st.button("⬅️ 상위"):
            st.session_state.current_path = os.path.dirname(st.session_state.current_path)

        try:
            items = sorted(os.listdir(st.session_state.current_path))
        except PermissionError:
            st.error(f"접근 권한이 없습니다: {st.session_state.current_path}")
            items = []
        except FileNotFoundError:
            st.error(f"디렉토리가 존재하지 않습니다: {st.session_state.current_path}")
            st.session_state.current_path = os.getcwd()
            items = sorted(os.listdir(st.session_state.current_path))

        for item in items:
            if item.startswith('.'): continue
            full_path = os.path.join(st.session_state.current_path, item)
            cl, cr = st.columns([4, 1])
            if os.path.isdir(full_path):
                if cl.button(f"📁 [{item}]", key=f"d_{item}"):
                    st.session_state.current_path = full_path
            else:
                ext = os.path.splitext(item)[1].lower()
                if ext in ['.pdf', '.epub']:
                    if cl.button(f"📄 {item}", key=f"f_{item}"):
                        import subprocess
                        import shutil
                        import platform

                        current_os = platform.system()

                        if current_os == "Windows":
                            sumatra_path = shutil.which('sumatrapdf')

                            if sumatra_path:
                                subprocess.Popen([sumatra_path, full_path])
                            else:
                                default_paths = [
                                    os.path.join("C:", "Program Files", "SumatraPDF", "SumatraPDF.exe"),
                                    os.path.join("C:", "Program Files (x86)", "SumatraPDF", "SumatraPDF.exe")
                                ]

                                found = False
                                for path in default_paths:
                                    if os.path.exists(path):
                                        subprocess.Popen([path, full_path])
                                        found = True
                                        break

                                if not found:
                                    try:
                                        os.startfile(full_path)
                                    except:
                                        st.error("SumatraPDF가 설치되지 않았습니다. PDF를 열 수 없습니다.")
                        elif current_os == "Darwin":
                            try:
                                subprocess.Popen(["open", full_path])
                            except:
                                st.error("PDF를 열 수 없습니다. 기본 뷰어를 확인하세요.")
                        elif current_os == "Linux":
                            try:
                                subprocess.Popen(["xdg-open", full_path])
                            except:
                                st.error("PDF를 열 수 없습니다. 기본 뷰어를 확인하세요.")
                        else:
                            st.error(f"지원하지 않는 운영체제: {current_os}")
                else:
                    if cl.button(f"📄 {item}", key=f"f_{item}"):
                        file_content = read_file(full_path)
                        st.session_state.v_content = f"#### 📄 [{item}]\n{file_content}"
                if cr.button("🧺", key=f"b_{item}"):
                    txt = read_file(full_path)
                    content_lines = txt.split('\n', 1)
                    if len(content_lines) >= 2 and content_lines[0].startswith('#### ['):
                        bible_ref = content_lines[0].replace('#### ', '')
                        st.session_state.basket.append({"file": item, "content": f"#### 📄 [{item}]\n{bible_ref}\n{content_lines[1]}"})
                    else:
                        st.session_state.basket.append({"file": item, "content": f"#### 📄 [{item}]\n{txt}"})
                    st.toast("담기 완료!")

    with t3:
        if st.session_state.basket:
            all_text = "\n\n".join([i['content'] for i in st.session_state.basket])
            st.code(all_text, language="text")

    st.divider()

    if st.session_state.scan_res:
        if st.button("📥 일괄 바구니담기", use_container_width=True):
            for res in st.session_state.scan_res:
                st.session_state.basket.append(res)
            st.toast(f"모든 결과 {len(st.session_state.scan_res)}개를 바구니에 담았습니다!")
            st.rerun()  # [버그수정] 바구니 숫자 즉시 갱신

    for i, res in enumerate(st.session_state.scan_res):
        cb, ca, cc, cd = st.columns([2.5, 1, 1, 1])
        if cb.button(f"📍 {res['file']}", key=f"res_{i}", use_container_width=True):
            if res['file'].startswith('📚'):
                st.session_state.v_content = f"#### 📄 [{res['file']}]\n{res['content']}"
            else:
                content_lines = res['content'].split('\n', 1)
                if len(content_lines) >= 2 and content_lines[0].startswith('#### ['):
                    bible_ref = content_lines[0].replace('#### ', '')
                    st.session_state.v_content = f"#### 📄 [{res['file']}]\n{bible_ref}\n{content_lines[1]}"
                else:
                    st.session_state.v_content = f"#### 📄 [{res['file']}]\n{res['content']}"
        if ca.button("🧺", key=f"ad_{i}"):
            st.session_state.basket.append(res)
            st.toast("바구니 저장!")
            st.rerun()  # [버그수정] 바구니 숫자 즉시 갱신

        # 카드형 보기 버튼 수정 - 별도 창으로 열기
        if cc.button("🔍", key=f"win_{i}"):
            content_to_display = ""
            if res['file'].startswith('📚'):
                content_to_display = f"#### 📄 [{res['file']}]\n{res['content']}"
            else:
                content_lines = res['content'].split('\n', 1)
                if len(content_lines) >= 2 and content_lines[0].startswith('#### ['):
                    bible_ref = content_lines[0].replace('#### ', '')
                    content_to_display = f"#### 📄 [{res['file']}]\n{bible_ref}\n{content_lines[1]}"
                else:
                    content_to_display = f"#### 📄 [{res['file']}]\n{res['content']}"

            content_escaped_lt = content_to_display.replace('<', '&lt;')
            content_escaped_gt = content_escaped_lt.replace('>', '&gt;')
            content_escaped_final = content_escaped_gt.replace('\n', '<br>')

            # 다크모드 HTML 생성
            html_content = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <title>BibleAI - Separate Viewer</title>
                <meta charset="UTF-8">
                <style>
                    body {{
                        font-family: Arial, sans-serif;
                        margin: 20px;
                        background-color: #1e1e1e;
                        color: #d4d4d4;
                    }}
                    .content {{
                        background-color: #2d2d30;
                        padding: 20px;
                        border-radius: 8px;
                        box-shadow: 0 2px 10px rgba(0,0,0,0.3);
                        line-height: 1.6;
                        white-space: pre-wrap;
                    }}
                    .controls {{
                        margin-bottom: 20px;
                    }}
                    button {{
                        background-color: #007acc;
                        color: white;
                        padding: 10px 15px;
                        border: none;
                        border-radius: 4px;
                        cursor: pointer;
                        margin-right: 10px;
                    }}
                    button:hover {{
                        background-color: #005a9e;
                    }}
                </style>
            </head>
            <body>
                <div class="controls">
                    <button onclick="window.print()">🖨️ 인쇄</button>
                    <button onclick="history.back()">🔙 뒤로가기</button>
                    <button onclick="window.close()">❌ 창 닫기</button>
                </div>
                <div class="content">
                    {content_escaped_final}
                </div>
            </body>
            </html>
            """

            with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix='.html', encoding='utf-8') as f:
                f.write(html_content)
                temp_file_path = f.name

            webbrowser.open('file://' + os.path.abspath(temp_file_path))

# ==============================
# ✅ [COL_R 영역 - 개편된 UI]
# ==============================

with col_r:
    st.markdown("### 🧠 대화형 에이전트 UI")

# --- [8. 우측 화면: 대화형 에이전트 UI] ---
    st.markdown("### 🤖 Bible Agent for Search & Preach")

    # 1. 상단: 결과 카드 뷰어 (i5-6400 시스템 보호를 위한 고정 높이 설정)
    if st.session_state.scan_res:
        st.markdown(f"#### 🔎 발견된 영적 보화 ({len(st.session_state.scan_res)}건)")
        with st.container(height=500):
            for i, res in enumerate(st.session_state.scan_res):
                with st.container(border=True):
                    # 파일명 옆에 아이콘 추가
                    st.markdown(f"**📄 {res['file']}**")

                    # 텍스트 미리보기 최적화 (메모리 절약)
                    content_preview = "\n".join(res['content'].split('\n')[:3]) + " ..." if len(res['content'].split('\n')) > 3 else res['content']
                    st.caption(content_preview)

                    btn_col1, btn_col2 = st.columns([1, 1])
                    with btn_col1:
                        if st.button(f"🧺 담기", key=f"add_basket_{i}", use_container_width=True):
                            if res not in st.session_state.basket:
                                st.session_state.basket.append(res)
                                st.toast(f"바구니에 담겼습니다!", icon="📥")
                                st.rerun()  # [버그수정] 바구니 숫자 즉시 갱신
                            else:
                                st.toast("이미 담겨 있습니다.", icon="⚠️")
                    with btn_col2:
                        # 카드형 보기 버튼도 별도 창으로 열기
                        if st.button(f"🔍 보기", key=f"view_detail_{i}", use_container_width=True):
                            content_to_display = f"#### 📄 [{res['file']}]\n{res['content']}"

                            content_escaped_lt = content_to_display.replace('<', '&lt;')
                            content_escaped_gt = content_escaped_lt.replace('>', '&gt;')
                            content_escaped_final = content_escaped_gt.replace('\n', '<br>')

                            # 다크모드 HTML 생성
                            html_content = f"""
                            <!DOCTYPE html>
                            <html>
                            <head>
                                <title>BibleAI - Separate Viewer</title>
                                <meta charset="UTF-8">
                                <style>
                                    body {{
                                        font-family: Arial, sans-serif;
                                        margin: 20px;
                                        background-color: #1e1e1e;
                                        color: #d4d4d4;
                                    }}
                                    .content {{
                                        background-color: #2d2d30;
                                        padding: 20px;
                                        border-radius: 8px;
                                        box-shadow: 0 2px 10px rgba(0,0,0,0.3);
                                        line-height: 1.6;
                                        white-space: pre-wrap;
                                    }}
                                    .controls {{
                                        margin-bottom: 20px;
                                    }}
                                    button {{
                                        background-color: #007acc;
                                        color: white;
                                        padding: 10px 15px;
                                        border: none;
                                        border-radius: 4px;
                                        cursor: pointer;
                                        margin-right: 10px;
                                    }}
                                    button:hover {{
                                        background-color: #005a9e;
                                    }}
                                </style>
                            </head>
                            <body>
                                <div class="controls">
                                    <button onclick="window.print()">🖨️ 인쇄</button>
                                    <button onclick="history.back()">🔙 뒤로가기</button>
                                    <button onclick="window.close()">❌ 창 닫기</button>
                                </div>
                                <div class="content">
                                    {content_escaped_final}
                                </div>
                            </body>
                            </html>
                            """

                            with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix='.html', encoding='utf-8') as f:
                                f.write(html_content)
                                temp_file_path = f.name

                            webbrowser.open('file://' + os.path.abspath(temp_file_path))

# --- 설치 가이드 ---
st.markdown("---")
st.subheader("📚 BibleAI 설치 가이드")
st.info("이 프로그램은 목회자와 선교사님의 사역을 돕기 위해 설계되었습니다. 설치 및 운영법은 아래 가이드를 참조하세요.")
_col1, _col2, _col3 = st.columns(3)
with _col1:
    if st.button("🪟 Windows 설치 가이드", key="guide_win"):
        webbrowser.open_new_tab("https://bonghgoo.tistory.com/569")
with _col2:
    if st.button("🍎 Mac 설치 가이드", key="guide_mac"):
        webbrowser.open_new_tab("https://bonghgoo.tistory.com/570?category=1300595")
with _col3:
    if st.button("🐧 Linux 설치 가이드", key="guide_linux"):
        webbrowser.open_new_tab("https://bonghgoo.tistory.com/571?category=1300595")

# 2. 중단: 지능형 명령어 입력창
st.markdown("---")
st.markdown("**💬 에이전트 명령어 입력/ 결과**")

# 검색 히스토리 트리거 처리
if 'trigger_search' in st.session_state and st.session_state.trigger_search:
    user_input = st.session_state.trigger_search
    st.session_state.trigger_search = None  # 트리거 초기화
else:
    user_input = st.chat_input("🔍 주제어 검색 / 명령(생성, 담아줘) / 로컬AI 질문(?)")

if user_input:
    # Store the current user input
    st.session_state.last_user_input = user_input
    # [A] 로컬 AI 샌드박스: '?'로 시작하면 무절 로컬 AI에게!
    if user_input.startswith("?"):
        question = user_input[1:].strip()
        with st.chat_message("user"):
            st.write(user_input)

        with st.chat_message("ai", avatar="🤖"):
            response_placeholder = st.empty()
            full_response = ""
            try:
                # --- [수정 시작] 이 부분이 엔진별로 답변을 가져오는 핵심입니다 ---
                if st.session_state.selected_engine == "groq":
                    response = client.chat.completions.create(
                        model=st.session_state.selected_model,
                        messages=[{'role': 'user', 'content': question}],
                        stream=True
                    )
                    for chunk in response:
                        if chunk.choices[0].delta.content:
                            full_response += chunk.choices[0].delta.content
                            response_placeholder.markdown(full_response + "▌")
                else:
                    response = ollama.chat(
                        model=st.session_state.selected_model,
                        messages=[{'role': 'user', 'content': question}],
                        stream=True
                    )
                    for chunk in response:
                        full_response += chunk['message']['content']
                        response_placeholder.markdown(full_response + "▌")
                # --- [수정 끝] 여기서부터 아래의 버튼 로직은 그대로 두시면 됩니다 ---
                
                # Display the final response
                response_placeholder.markdown(full_response)
                
                # Store the response in session state
                st.session_state.ai_response = full_response
                st.session_state.show_ai_response = True

                # (이후 사용자님의 복사 버튼, 보기 버튼 코드가 쭉 이어집니다...)

                # Add copy and basket buttons
                col1, col2 = st.columns([1, 3])
                with col1:
                    # New HTML/JS copy button that doesn't cause a rerun
                    unique_id = f"copy_new_{abs(hash(full_response))}"
                    escaped_text = json.dumps(full_response)
                    copy_html = f'''
                    <button id="{unique_id}" onclick="copyToClipboard_{unique_id}(this, event)">📋 응답 복사</button>
                    <script>
                        function copyToClipboard_{unique_id}(button, event) {{
                            event.preventDefault();
                            const text = {escaped_text};
                            navigator.clipboard.writeText(text).then(function() {{
                                const originalText = button.innerHTML;
                                button.innerHTML = '✅ 복사 완료!';
                                setTimeout(function() {{
                                    button.innerHTML = originalText;
                                }}, 2000);
                            }}, function(err) {{
                                console.error('Could not copy text: ', err);
                                button.innerHTML = '❌ 복사 실패';
                            }});
                        }}
                    </script>
                    '''
                    components.html(copy_html, height=40)

            except Exception as e:
                st.error(f"Grpq 연결을 확인해 주세요: {e}")

    # [B] 특수 명령어 처리: '생성'이 포함된 경우 (조립 및 클립보드 복사)
    elif "생성" in user_input:
        if not st.session_state.basket:
            st.warning("🧺 바구니가 비어 있습니다. 자료를 먼저 담아주세요.")
        else:
            # 바구니 내용을 텍스트로 결합
            combined_context = "\n\n".join([f"--- {item['file']} ---\n{item['content']}" for item in st.session_state.basket])
            final_prompt = get_custom_prompt(combined_context)

            # 클립보드 복사 및 출력
            try:
                pyperclip.copy(final_prompt)
                st.success("✅ 설교 프롬프트가 조립되어 클립보드에 복사되었습니다!")
            except:
                st.warning("⚠️ 클립보드 복사 실패. 아래 내용을 수동으로 복사하세요.")

            with st.expander("📝 생성된 프롬프트 확인", expanded=True):
                st.code(final_prompt, language="markdown")

    # [C] 데이터 축적: '담아줘'가 포함된 경우
    elif "담아줘" in user_input:
        # 검색 결과 모두 바구니에 담기
        if st.session_state.scan_res:
            added_count = 0
            for item in st.session_state.scan_res:
                if item not in st.session_state.basket:
                    st.session_state.basket.append(item)
                    # [NEW 3] 현재 그룹에도 추가
                    st.session_state.basket_groups[st.session_state.current_group].append(item)
                    added_count += 1

            if added_count > 0:
                st.success(f"🧺 총 {added_count}개의 검색 결과를 '{st.session_state.current_group}'에 담았습니다.")
                st.rerun()  # [버그수정] 바구니 숫자 즉시 갱신
            else:
                st.info("이미 모든 결과가 바구니에 담겨 있습니다.")
        else:
            st.warning("⚠️ 담을 검색 결과가 없습니다. 먼저 검색을 수행해 주세요.")

    # [D] 일반 검색 실행 (개선된 주제어 검색 로직)
    else:
        with st.chat_message("user"):
            st.write(user_input)

        with st.chat_message("ai"):
            # [버그수정] 폴더 미설정 시 경고하고 중단 (무한루프 방지)
            if not selected_folders:
                st.warning(
                    "⚠️ **검색 범위(폴더)가 설정되지 않았습니다.**\n\n"
                    "왼쪽 사이드바에서 '📍 검색 범위(폴더)'를 먼저 선택한 후 다시 시도해 주세요.\n\n"
                    "💡 *폴더를 선택하지 않으면 주석 모듈(DB)을 읽을 수 없어 검색이 올바르게 작동하지 않습니다.*"
                )
            else:
                # 1. 성경 및 파일 검색 수행
                bible_results = search_bible_sqlite(user_input)

                # 개선된 파일 검색 (조건 검색 + 컨텍스트 기반 스니펫)
                file_results = search_files_advanced(
                    user_input,
                    selected_folders
                )

                # [버그수정] 주석 모듈(DB) 검색 병행 실행 - 기존에는 누락되었음
                # 현재 검색어를 성경 절로 해석하기보다 텍스트 포함 여부로 DB 주석을 별도 탐색
                db_results = []
                try:
                    com_files = scan_commentary_files(selected_folders)
                    if com_files:
                        # 주석 DB에서 텍스트 포함 검색 (키워드 매칭)
                        query_lower = user_input.lower()
                        for com_path in com_files:
                            try:
                                conn = sqlite3.connect(com_path)
                                cur = conn.cursor()
                                # 가능한 테이블/컬럼 조합 시도
                                for table, col in [
                                    ("commentary", "data"), ("content", "data"),
                                    ("commentaries", "text"), ("Verses", "Comments"),
                                    ("VerseCommentary", "Comments"),
                                ]:
                                    try:
                                        cur.execute(
                                            f"SELECT {col} FROM {table} WHERE lower({col}) LIKE ? LIMIT 5",
                                            (f"%{query_lower}%",)
                                        )
                                        rows = cur.fetchall()
                                        for row in rows:
                                            if row[0]:
                                                decoded = decode_rtf(row[0]) if isinstance(row[0], (str, bytes)) else str(row[0])
                                                snippet = decoded.strip()[:500]
                                                if snippet:
                                                    fname = os.path.basename(com_path)
                                                    db_results.append({
                                                        "file": f"📚 {fname}",
                                                        "content": f"#### 📚 [{fname}]\n{snippet}",
                                                        "relevance_score": snippet.lower().count(query_lower) * 10
                                                    })
                                        break  # 성공한 테이블에서만 읽음
                                    except Exception:
                                        continue
                                conn.close()
                            except Exception:
                                pass
                except Exception:
                    pass

                total_results = bible_results + file_results + db_results

            if not selected_folders:
                pass  # 위에서 이미 경고 표시
            elif not locals().get('total_results'):
                st.info(f"🔍 '{user_input}' 검색 결과가 없습니다.")
            else:
                total_results = locals().get('total_results', [])
                st.session_state.last_search_results = total_results
                st.session_state.scan_res = total_results

                # [NEW 2] 검색 히스토리에 저장
                save_search_history(user_input)

                st.subheader(f"🔎 에이전트 검색 결과 ({len(total_results)}건)")
                st.caption(f"검색어: **{user_input}** | 조건 검색(+필수, -제외) 지원")

                # --- 결과 출력 시작 ---
                for i, item in enumerate(total_results):
                    with st.container(border=True):
                        # [NEW 1] 관련도 점수 표시
                        score = item.get('relevance_score', 0)
                        if score > 0:
                            st.markdown(f"### 📂 {item['file']} 🎯 관련도: {score}점")
                        else:
                            st.markdown(f"### 📂 {item['file']}")

                        # 미리보기: 처음 300자만 표시
                        content_lines = item['content'].split('\n')
                        preview_lines = []
                        char_count = 0
                        for line in content_lines:
                            if char_count + len(line) > 300:
                                preview_lines.append(line[:300-char_count] + "...")
                                break
                            preview_lines.append(line)
                            char_count += len(line)

                        preview = "\n".join(preview_lines)
                        st.write(preview)

                        # [다크모드 전문 보기 버튼 로직 - 수정된 버전]
                        safe_content = item['content'].replace("'", "\'").replace("\n", "\n").replace('"', '"')

                        # JavaScript에서 안전하게 사용할 수 있도록 이스케이프 처리
                        html_safe_content = item['content'].replace('<', '&lt;').replace('>', '&gt;').replace('\n', '<br>')

                        viewer_html = f"""
                        <div style="margin-top: 10px;">
                            <button onclick="openDarkViewer{i}()"
                               style="cursor:pointer; background-color:#1e1e1e; color:#569cd6; padding:8px 20px; border-radius:5px; border:1px solid #569cd6; font-size:14px; font-weight:bold;">🔍 다크모드 전문 보기</button>
                        </div>
                        <script>
                        function openDarkViewer{i}() {{
                            var win = window.open("", "_blank", "width=900,height=850");
                            var content = `{html_safe_content}`;
                            win.document.write(`
                                <html>
                                <head>
                                    <title>Dark Viewer - {item['file']}</title>
                                    <meta charset="UTF-8">
                                    <style>
                                        body {{ background: #1e1e1e; color: #d4d4d4; font-family: 'Malgun Gothic', sans-serif; padding: 40px; line-height: 1.9; }}
                                        .header {{ border-bottom: 2px solid #3e3e42; padding-bottom: 15px; margin-bottom: 25px; display: flex; justify-content: space-between; align-items: center; }}
                                        h2 {{ color: #569cd6; margin: 0; font-size: 1.5em; }}
                                        .content {{ white-space: pre-wrap; font-size: 1.15em; letter-spacing: 0.05em; }}
                                        .no-print {{ background: #333; color: #fff; border: none; padding: 7px 15px; cursor: pointer; border-radius: 4px; }}
                                        @media print {{ .no-print {{ display: none; }} body {{ background: white; color: black; }} }}
                                    </style>
                                </head>
                                <body>
                                    <div class="header">
                                        <h2>📂 {item['file']}</h2>
                                        <button class="no-print" onclick="window.print()">🖨️ 프린트</button>
                                    </div>
                                    <div class="content">` + content + `</div>
                                    <br><br>
                                    <center><button class="no-print" onclick="window.close()" style="background:#444;">닫기</button></center>
                                </body>
                                </html>
                            `);
                        }}
                        </script>
                        """
                        components.html(viewer_html, height=70)

                        # [NEW 4, 5] 복사 및 AI 요약 버튼 추가 - 이 부분 삭제 2026-02-26

    # Display stored AI response if it exists and no new input
    if st.session_state.show_ai_response and st.session_state.ai_response and not user_input:
        with st.chat_message("ai", avatar="🤖"):
            st.markdown(st.session_state.ai_response)
            
            # Add copy and basket buttons for the stored response
            col1, col2 = st.columns([1, 3])
            with col1:
                # New HTML/JS copy button
                unique_id = f"copy_stored_{abs(hash(st.session_state.ai_response))}"
                escaped_text = json.dumps(st.session_state.ai_response)
                copy_html = f'''
                <button id="{unique_id}" onclick="copyToClipboard_{unique_id}(this, event)">📋 응답 복사</button>
                <script>
                    function copyToClipboard_{unique_id}(button, event) {{
                        event.preventDefault();
                        const text = {escaped_text};
                        navigator.clipboard.writeText(text).then(function() {{
                            const originalText = button.innerHTML;
                            button.innerHTML = '✅ 복사 완료!';
                            setTimeout(function() {{
                                button.innerHTML = originalText;
                            }}, 2000);
                        }}, function(err) {{
                            console.error('Could not copy text: ', err);
                            button.innerHTML = '❌ 복사 실패';
                        }});
                    }}
                </script>
                '''
                components.html(copy_html, height=40)
    
    

    # 하단 제어바 제거됨

st.markdown("**제작: 경인노회 (<a href='https://kinohoi.blogspot.com' target='_blank'>https://kinohoi.blogspot.com</a>) 신학연구원 BibleAI Team**", unsafe_allow_html=True)
